"""DOCX 解析器 —— 使用 python-docx 提取 Word 文档中的文本。"""

from docx import Document

from backend.infrastructure.parsers.base import (
    BLOCK_GENERIC,
    BLOCK_HEADING,
    BLOCK_PARAGRAPH,
    ParsedResumeText,
    ResumeParser,
    TextBlock,
)


class DocxResumeParser(ResumeParser):
    """Word 文档解析器，提取段落文本和表格内容。"""

    @property
    def version(self) -> str:
        return "docx-python-docx-v1"

    def parse(self, file_path: str) -> ParsedResumeText:
        doc = Document(file_path)

        parts: list[str] = []
        blocks: list[TextBlock] = []

        # 提取所有段落文本，利用 Word 样式名识别标题
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            parts.append(text)
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            block_type = BLOCK_HEADING if "heading" in style_name or "title" in style_name else BLOCK_PARAGRAPH
            blocks.append(TextBlock(type=block_type, text=text))

        # 提取表格内容（简历中常用表格排版）
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    row_text = " | ".join(cells)
                    parts.append(row_text)
                    blocks.append(TextBlock(type=BLOCK_GENERIC, text=row_text))

        return ParsedResumeText(
            raw_text="\n".join(parts),
            page_count=None,  # DOCX 格式无法直接获取页数
            blocks=blocks,
        )
